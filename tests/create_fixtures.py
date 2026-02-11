"""Generate .docx test fixture files for the Word handler tests."""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


FIXTURES_DIR = Path(__file__).parent / "fixtures"
FIXTURES_DIR.mkdir(exist_ok=True)


def create_table_questionnaire() -> None:
    """Fixture 1: two-column table questionnaire (question | answer).

    Simulates a typical legal/insurance questionnaire with a header row
    and several question rows where the right column is empty (to be filled).
    """
    doc = Document()

    # Title
    title = doc.add_heading("Vendor Due Diligence Questionnaire", level=1)

    # Intro paragraph
    doc.add_paragraph(
        "Please complete all fields below. Attach additional documentation where indicated."
    )

    # Create the questionnaire table: 2 columns, header + 6 question rows
    table = doc.add_table(rows=7, cols=2, style="Table Grid")

    # Header row
    hdr = table.rows[0]
    hdr.cells[0].text = "Question"
    hdr.cells[1].text = "Answer"
    for cell in hdr.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Calibri"

    # Questions
    questions = [
        "What is the full legal name of your company?",
        "What is your company's principal address?",
        "What is the name and title of the primary contact?",
        "Does your company maintain cyber liability insurance? If so, state the coverage limit.",
        "Describe your company's data privacy and protection policies.",
        "List any regulatory certifications your company holds (e.g., SOC 2, ISO 27001).",
    ]

    for i, q in enumerate(questions, start=1):
        row = table.rows[i]
        # Question cell - set formatting
        q_para = row.cells[0].paragraphs[0]
        q_run = q_para.add_run(q)
        q_run.font.name = "Calibri"
        q_run.font.size = Pt(10)

        # Answer cell - leave empty (the AI/agent will fill this)
        a_para = row.cells[1].paragraphs[0]
        a_run = a_para.add_run("")
        a_run.font.name = "Calibri"
        a_run.font.size = Pt(10)

    # Second table: a smaller one with different formatting
    doc.add_paragraph("")  # spacer
    doc.add_heading("Section 2: Financial Information", level=2)

    table2 = doc.add_table(rows=4, cols=2, style="Table Grid")
    hdr2 = table2.rows[0]
    hdr2.cells[0].text = "Item"
    hdr2.cells[1].text = "Details"
    for cell in hdr2.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Arial"

    financial_questions = [
        "Annual revenue (most recent fiscal year):",
        "Total number of employees:",
        "Year of incorporation:",
    ]

    for i, q in enumerate(financial_questions, start=1):
        row = table2.rows[i]
        q_para = row.cells[0].paragraphs[0]
        q_run = q_para.add_run(q)
        q_run.font.name = "Arial"
        q_run.font.size = Pt(10)

        # Leave answer cell empty
        a_para = row.cells[1].paragraphs[0]
        a_run = a_para.add_run("")
        a_run.font.name = "Arial"
        a_run.font.size = Pt(10)

    out_path = FIXTURES_DIR / "table_questionnaire.docx"
    doc.save(str(out_path))
    print(f"Created {out_path}")


def create_placeholder_form() -> None:
    """Fixture 2: paragraph form with placeholder text.

    Simulates a contract/form document where answers are inline placeholders
    like '[Enter here]', '___', etc.
    """
    doc = Document()

    doc.add_heading("Non-Disclosure Agreement", level=1)

    # Paragraph with inline placeholder
    p1 = doc.add_paragraph()
    r1 = p1.add_run("This Non-Disclosure Agreement (the \"Agreement\") is entered into as of ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r_date = p1.add_run("[Enter date]")
    r_date.font.name = "Times New Roman"
    r_date.font.size = Pt(12)
    r_date.italic = True
    r1b = p1.add_run(" by and between:")
    r1b.font.name = "Times New Roman"
    r1b.font.size = Pt(12)

    # Party 1
    p2 = doc.add_paragraph()
    r2a = p2.add_run("Party 1 (\"Disclosing Party\"): ")
    r2a.font.name = "Times New Roman"
    r2a.font.size = Pt(12)
    r2a.bold = True
    r2b = p2.add_run("_______________")
    r2b.font.name = "Times New Roman"
    r2b.font.size = Pt(12)

    # Party 1 address
    p2addr = doc.add_paragraph()
    r2addr_label = p2addr.add_run("Address: ")
    r2addr_label.font.name = "Times New Roman"
    r2addr_label.font.size = Pt(12)
    r2addr_val = p2addr.add_run("[Enter here]")
    r2addr_val.font.name = "Times New Roman"
    r2addr_val.font.size = Pt(12)
    r2addr_val.italic = True

    # Party 2
    p3 = doc.add_paragraph()
    r3a = p3.add_run("Party 2 (\"Receiving Party\"): ")
    r3a.font.name = "Times New Roman"
    r3a.font.size = Pt(12)
    r3a.bold = True
    r3b = p3.add_run("_______________")
    r3b.font.name = "Times New Roman"
    r3b.font.size = Pt(12)

    # Party 2 address
    p3addr = doc.add_paragraph()
    r3addr_label = p3addr.add_run("Address: ")
    r3addr_label.font.name = "Times New Roman"
    r3addr_label.font.size = Pt(12)
    r3addr_val = p3addr.add_run("[Enter here]")
    r3addr_val.font.name = "Times New Roman"
    r3addr_val.font.size = Pt(12)
    r3addr_val.italic = True

    # Body paragraph
    doc.add_paragraph("")
    body = doc.add_paragraph()
    rb1 = body.add_run(
        "The parties agree that any confidential information disclosed under this "
        "Agreement shall be protected for a period of "
    )
    rb1.font.name = "Times New Roman"
    rb1.font.size = Pt(12)
    rb2 = body.add_run("[Enter number]")
    rb2.font.name = "Times New Roman"
    rb2.font.size = Pt(12)
    rb2.italic = True
    rb3 = body.add_run(" years from the date of disclosure.")
    rb3.font.name = "Times New Roman"
    rb3.font.size = Pt(12)

    # Governing law
    law = doc.add_paragraph()
    rl1 = law.add_run("Governing Law: This Agreement shall be governed by the laws of ")
    rl1.font.name = "Times New Roman"
    rl1.font.size = Pt(12)
    rl2 = law.add_run("[Enter jurisdiction]")
    rl2.font.name = "Times New Roman"
    rl2.font.size = Pt(12)
    rl2.italic = True
    rl3 = law.add_run(".")
    rl3.font.name = "Times New Roman"
    rl3.font.size = Pt(12)

    # Signature block
    doc.add_paragraph("")
    doc.add_heading("Signatures", level=2)

    sig1 = doc.add_paragraph()
    s1a = sig1.add_run("Disclosing Party Signature: ")
    s1a.font.name = "Times New Roman"
    s1a.font.size = Pt(12)
    s1b = sig1.add_run("_______________")
    s1b.font.name = "Times New Roman"
    s1b.font.size = Pt(12)

    sig1_name = doc.add_paragraph()
    s1n_label = sig1_name.add_run("Printed Name: ")
    s1n_label.font.name = "Times New Roman"
    s1n_label.font.size = Pt(12)
    s1n_val = sig1_name.add_run("[Enter here]")
    s1n_val.font.name = "Times New Roman"
    s1n_val.font.size = Pt(12)
    s1n_val.italic = True

    sig2 = doc.add_paragraph()
    s2a = sig2.add_run("Receiving Party Signature: ")
    s2a.font.name = "Times New Roman"
    s2a.font.size = Pt(12)
    s2b = sig2.add_run("_______________")
    s2b.font.name = "Times New Roman"
    s2b.font.size = Pt(12)

    sig2_name = doc.add_paragraph()
    s2n_label = sig2_name.add_run("Printed Name: ")
    s2n_label.font.name = "Times New Roman"
    s2n_label.font.size = Pt(12)
    s2n_val = sig2_name.add_run("[Enter here]")
    s2n_val.font.name = "Times New Roman"
    s2n_val.font.size = Pt(12)
    s2n_val.italic = True

    out_path = FIXTURES_DIR / "placeholder_form.docx"
    doc.save(str(out_path))
    print(f"Created {out_path}")


if __name__ == "__main__":
    create_table_questionnaire()
    create_placeholder_form()
